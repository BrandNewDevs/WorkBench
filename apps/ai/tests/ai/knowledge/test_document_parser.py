"""Tests for parsing only backend-supplied knowledge documents."""

from io import BytesIO
from pathlib import Path
from typing import cast
from zipfile import ZIP_DEFLATED, ZipFile

import pymupdf
import pytest
from docx import Document as create_docx

from app.ai.errors import (
    CorruptKnowledgeInput,
    KnowledgeInputTooLarge,
    UnsupportedKnowledgeInput,
)
from app.ai.knowledge.config import KnowledgeProcessingSettings
from app.ai.knowledge.document_parser import (
    DOCX_MIME,
    PDF_MIME,
    TEXT_MIME,
    LocalDocumentParser,
)
from app.ai.schemas import ApprovedPath, SourceDocument


def pdf_bytes() -> bytes:
    """Create a small sanitized two-page SOP fixture."""

    with pymupdf.open() as document:  # type: ignore[no-untyped-call]
        first_page = document.new_page(width=595, height=842)
        first_page.insert_text((72, 72), "ISOLATION PROCEDURE", fontsize=16)
        first_page.insert_text((72, 110), "Close the inlet valve before inspection.")
        second_page = document.new_page(width=595, height=842)
        second_page.insert_text((72, 72), "RECORDS", fontsize=16)
        second_page.insert_text((72, 110), "Record the observed valve condition.")
        return cast(bytes, document.tobytes())


def docx_bytes() -> bytes:
    """Create a sanitized prior-note fixture with native heading styles."""

    document = create_docx()
    document.add_heading("Previous decision", level=1)
    document.add_paragraph("Thickness measurement was approved before repair.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Owner"
    table.cell(0, 1).text = "Inspection team"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def source_document(
    content: bytes,
    mime_type: str,
    *,
    document_id: str = "document-1",
    document_name: str = "document",
) -> SourceDocument:
    """Wrap sanitized bytes in application-controlled metadata."""

    return SourceDocument(
        document_id=document_id,
        document_name=document_name,
        mime_type=mime_type,
        source_id=f"source-{document_id}",
        content=content,
    )


@pytest.mark.parametrize(
    ("document", "expected_pages", "expected_heading"),
    [
        (source_document(pdf_bytes(), PDF_MIME), {1, 2}, "ISOLATION PROCEDURE"),
        (source_document(docx_bytes(), DOCX_MIME), {1}, "Previous decision"),
        (
            source_document(
                b"APPROVAL NOTE TEMPLATE\n\nSubject\n\nRecommendation",
                TEXT_MIME,
            ),
            {1},
            "APPROVAL NOTE TEMPLATE",
        ),
    ],
)
def test_supported_documents_preserve_page_and_heading_structure(
    document: SourceDocument,
    expected_pages: set[int],
    expected_heading: str,
) -> None:
    """Parse the three MVP corpus document forms without a model call."""

    parsed = LocalDocumentParser().parse(document)

    assert {block.page_number for block in parsed.blocks} == expected_pages
    assert any(block.is_heading and block.text == expected_heading for block in parsed.blocks)
    assert parsed.source_id == document.effective_source_id


def test_plain_text_title_case_headings_create_section_boundaries() -> None:
    """Recognize common government-note headings without Markdown markers."""

    parsed = LocalDocumentParser().parse(
        source_document(
            b"Subject\n\nRecord the finding\n\n"
            b"Recommendation\n\nApprove measurement.",
            TEXT_MIME,
        )
    )

    headings = [block.text for block in parsed.blocks if block.is_heading]
    assert headings == ["Subject", "Recommendation"]
    assert any(block.text == "Record the finding" for block in parsed.blocks)


def test_parser_reads_the_exact_approved_path(tmp_path: Path) -> None:
    """Accept Backend 2's wrapper without discovering adjacent files."""

    approved_file = tmp_path / "approved-note.txt"
    approved_file.write_text("DECISION\n\nApprove inspection.", encoding="utf-8")
    document = SourceDocument(
        document_id="approved-note",
        document_name="approved-note.txt",
        mime_type=TEXT_MIME,
        approved_path=ApprovedPath(
            path=approved_file,
            source_id="approved-source",
            session_id="session-1",
        ),
    )

    parsed = LocalDocumentParser().parse(document)

    assert parsed.source_id == "approved-source"
    assert any(block.text == "Approve inspection." for block in parsed.blocks)


@pytest.mark.parametrize(
    ("document", "expected_error"),
    [
        (source_document(b"binary", "application/octet-stream"), UnsupportedKnowledgeInput),
        (source_document(b"not-a-pdf", PDF_MIME), CorruptKnowledgeInput),
        (source_document(b"\xff\xfe", TEXT_MIME), CorruptKnowledgeInput),
    ],
)
def test_unsupported_or_corrupt_documents_return_typed_errors(
    document: SourceDocument,
    expected_error: type[UnsupportedKnowledgeInput],
) -> None:
    """Hide parser-specific failures behind stable AI error types."""

    with pytest.raises(expected_error):
        LocalDocumentParser().parse(document)


def test_parser_enforces_input_size_limit() -> None:
    """Reject oversized content before document parsing allocates more memory."""

    parser = LocalDocumentParser(KnowledgeProcessingSettings(max_input_bytes=8))

    with pytest.raises(KnowledgeInputTooLarge):
        parser.parse(source_document(b"more than eight bytes", TEXT_MIME))


def test_docx_archive_expansion_is_bounded_before_parsing() -> None:
    """Reject a compressed archive whose expanded content exceeds the safe limit."""

    output = BytesIO()
    with ZipFile(output, mode="w", compression=ZIP_DEFLATED) as archive:
        archive.writestr("word/document.xml", b"x" * 2_000)
    parser = LocalDocumentParser(
        KnowledgeProcessingSettings(max_docx_uncompressed_bytes=1_000)
    )

    with pytest.raises(KnowledgeInputTooLarge, match="decompressed byte limit"):
        parser.parse(source_document(output.getvalue(), DOCX_MIME))

"""Local parsers for approved PDF, DOCX, text, and Markdown documents."""

import re
from io import BytesIO
from zipfile import BadZipFile, ZipFile

import pymupdf
from docx import Document as open_docx
from docx.opc.exceptions import PackageNotFoundError

from app.ai.errors import (
    CorruptKnowledgeInput,
    KnowledgeInputTooLarge,
    UnsupportedKnowledgeInput,
)
from app.ai.knowledge.config import KnowledgeProcessingSettings
from app.ai.knowledge.contracts import DocumentBlock, ParsedDocument
from app.ai.schemas import SourceDocument

PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
TEXT_MIME = "text/plain"
MARKDOWN_MIME = "text/markdown"
SUPPORTED_KNOWLEDGE_MIME_TYPES = frozenset(
    {PDF_MIME, DOCX_MIME, TEXT_MIME, MARKDOWN_MIME}
)

_PDF_ERRORS = (
    pymupdf.EmptyFileError,
    pymupdf.FileDataError,
    pymupdf.mupdf.FzErrorBase,
    RuntimeError,
    ValueError,
    OSError,
)
_MARKDOWN_HEADING = re.compile(r"^#{1,6}\s+(?P<title>.+)$")
_SUBSECTION_HEADING = re.compile(r"^\d+(?:\.\d+)+(?:\.?\s+)(?P<title>\S.+)$")


class LocalDocumentParser:
    """Read only the supplied document and expose its page/heading structure."""

    def __init__(self, settings: KnowledgeProcessingSettings | None = None) -> None:
        self._settings = settings or KnowledgeProcessingSettings()

    def parse(self, document: SourceDocument) -> ParsedDocument:
        """Decode one supported document without following neighboring paths."""

        data = self._read_document(document)
        mime_type = document.mime_type.partition(";")[0].strip().lower()
        if mime_type not in SUPPORTED_KNOWLEDGE_MIME_TYPES:
            raise UnsupportedKnowledgeInput(f"unsupported knowledge MIME type: {mime_type}")

        if mime_type == PDF_MIME:
            blocks = self._parse_pdf(data)
        elif mime_type == DOCX_MIME:
            blocks = self._parse_docx(data)
        else:
            blocks = self._parse_text(data)

        if not blocks:
            raise CorruptKnowledgeInput("knowledge document contains no extractable text")
        if len(blocks) > self._settings.max_document_blocks:
            raise KnowledgeInputTooLarge("knowledge document exceeds the block limit")
        return ParsedDocument(
            source_id=document.effective_source_id,
            document_id=document.document_id,
            document_name=document.document_name,
            mime_type=mime_type,
            blocks=blocks,
        )

    def _read_document(self, document: SourceDocument) -> bytes:
        if document.content is not None:
            self._require_bounded_size(len(document.content))
            return document.content

        approved = document.approved_path
        if approved is None:  # pragma: no cover - protected by schema validation
            raise CorruptKnowledgeInput("knowledge document has no approved input")
        try:
            size = approved.path.stat().st_size
            if not approved.path.is_file():
                raise CorruptKnowledgeInput("approved knowledge input is not a regular file")
            self._require_bounded_size(size)
            return approved.path.read_bytes()
        except UnsupportedKnowledgeInput:
            raise
        except OSError as error:
            raise CorruptKnowledgeInput("approved knowledge input could not be read") from error

    def _require_bounded_size(self, size: int) -> None:
        if size <= 0:
            raise CorruptKnowledgeInput("knowledge document is empty")
        if size > self._settings.max_input_bytes:
            raise KnowledgeInputTooLarge("knowledge document exceeds the byte limit")

    def _parse_pdf(self, data: bytes) -> tuple[DocumentBlock, ...]:
        try:
            with pymupdf.open(  # type: ignore[no-untyped-call]
                stream=data, filetype="pdf"
            ) as pdf:
                if pdf.needs_pass:
                    raise CorruptKnowledgeInput("encrypted knowledge PDFs are not supported")
                if pdf.page_count == 0:
                    raise CorruptKnowledgeInput("knowledge PDF contains no pages")
                if pdf.page_count > self._settings.max_pdf_pages:
                    raise KnowledgeInputTooLarge("knowledge PDF exceeds the page limit")

                blocks: list[DocumentBlock] = []
                for page_index in range(pdf.page_count):
                    page = pdf.load_page(page_index)
                    raw_blocks = page.get_text("blocks", sort=True)
                    for raw_block in raw_blocks:
                        if len(raw_block) < 7 or raw_block[6] != 0:
                            continue
                        blocks.extend(
                            self._blocks_from_text(str(raw_block[4]), page_index + 1)
                        )
                return tuple(blocks)
        except UnsupportedKnowledgeInput:
            raise
        except _PDF_ERRORS as error:
            raise CorruptKnowledgeInput("knowledge PDF could not be decoded safely") from error

    def _parse_docx(self, data: bytes) -> tuple[DocumentBlock, ...]:
        self._validate_docx_archive(data)
        try:
            document = open_docx(BytesIO(data))
        except (PackageNotFoundError, BadZipFile, KeyError, ValueError, OSError) as error:
            raise CorruptKnowledgeInput("knowledge DOCX could not be decoded safely") from error

        blocks: list[DocumentBlock] = []
        for paragraph in document.paragraphs:
            text = self._normalize_paragraph(paragraph.text)
            if not text:
                continue
            style_name = paragraph.style.name if paragraph.style is not None else ""
            blocks.append(
                DocumentBlock(
                    # Native DOCX has no stable page map; MVP corpus DOCX files are one page.
                    page_number=1,
                    text=self._strip_heading_marker(text),
                    is_heading=style_name.lower().startswith("heading")
                    or self._looks_like_heading(text),
                )
            )

        for table in document.tables:
            for row in table.rows:
                cells = [self._normalize_paragraph(cell.text) for cell in row.cells]
                text = " | ".join(cell for cell in cells if cell)
                if text:
                    blocks.append(DocumentBlock(page_number=1, text=text))
        return tuple(blocks)

    def _validate_docx_archive(self, data: bytes) -> None:
        try:
            with ZipFile(BytesIO(data)) as archive:
                entries = archive.infolist()
        except (BadZipFile, OSError) as error:
            raise CorruptKnowledgeInput("knowledge DOCX is not a valid archive") from error

        if len(entries) > self._settings.max_docx_entries:
            raise KnowledgeInputTooLarge("knowledge DOCX exceeds the archive entry limit")
        if any(entry.flag_bits & 0x1 for entry in entries):
            raise CorruptKnowledgeInput("encrypted knowledge DOCX files are not supported")
        uncompressed_size = sum(entry.file_size for entry in entries)
        if uncompressed_size > self._settings.max_docx_uncompressed_bytes:
            raise KnowledgeInputTooLarge("knowledge DOCX exceeds the decompressed byte limit")

    def _parse_text(self, data: bytes) -> tuple[DocumentBlock, ...]:
        try:
            decoded = data.decode("utf-8-sig")
        except UnicodeDecodeError as error:
            raise CorruptKnowledgeInput("knowledge text must use UTF-8 encoding") from error

        blocks: list[DocumentBlock] = []
        for page_number, page_text in enumerate(decoded.split("\f"), start=1):
            for paragraph in re.split(r"\n\s*\n", page_text):
                blocks.extend(self._blocks_from_text(paragraph, page_number))
        return tuple(blocks)

    def _blocks_from_text(self, value: str, page_number: int) -> list[DocumentBlock]:
        text = self._normalize_paragraph(value)
        if not text:
            return []

        lines = text.splitlines()
        first_line = lines[0]
        if len(lines) > 1 and self._looks_like_heading(first_line):
            heading = DocumentBlock(
                page_number=page_number,
                text=self._strip_heading_marker(first_line),
                is_heading=True,
            )
            remainder = self._normalize_paragraph("\n".join(lines[1:]))
            return [heading, DocumentBlock(page_number=page_number, text=remainder)]

        return [
            DocumentBlock(
                page_number=page_number,
                text=self._strip_heading_marker(text),
                is_heading=self._looks_like_heading(text),
            )
        ]

    def _looks_like_heading(self, text: str) -> bool:
        collapsed = " ".join(text.split())
        if not collapsed or len(collapsed) > self._settings.max_heading_chars:
            return False
        if len(text.splitlines()) != 1:
            return False
        if _MARKDOWN_HEADING.fullmatch(collapsed) or _SUBSECTION_HEADING.fullmatch(collapsed):
            return True
        letters = [character for character in collapsed if character.isalpha()]
        if bool(letters) and collapsed.upper() == collapsed and len(collapsed.split()) <= 16:
            return True
        if re.match(r"^\d+[.)]\s", collapsed):
            return False
        has_sentence_ending = collapsed.endswith((".", "?", "!", ";"))
        return (
            not has_sentence_ending
            and (
                collapsed.endswith(":")
                or (len(collapsed.split()) <= 2 and collapsed.istitle())
            )
        )

    @staticmethod
    def _strip_heading_marker(text: str) -> str:
        markdown_match = _MARKDOWN_HEADING.fullmatch(text)
        if markdown_match is not None:
            return markdown_match.group("title").strip()
        return text.strip()

    @staticmethod
    def _normalize_paragraph(value: str) -> str:
        lines = [" ".join(line.split()) for line in value.replace("\r", "\n").splitlines()]
        return "\n".join(line for line in lines if line).strip()
